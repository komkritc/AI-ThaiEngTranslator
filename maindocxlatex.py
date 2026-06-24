#!/usr/bin/env python3
"""
translate_docx.py

Translates a Thai research proposal DOCX into publication-quality academic English
or vice versa using a single-stage LLM pipeline supporting Ollama, OpenRouter, and Typhoon.

Features:
- Modern 2026 GUI (CustomTkinter) with Dark/Light mode
- CLI fallback for automated pipelines
- Single-stage translation (direct translation)
- Supports Ollama (local), OpenRouter (cloud), and Typhoon (SCB 10X)
- Language direction: Thai->English or English->Thai
- API key persistence in .txt file
- Rate limiting with exponential backoff for Typhoon
- Auto-optimal chunk sizing based on content
- Time estimation and countdown
- Export to LaTeX (.tex)
- Robust API handling (Retries, Rate-limiting, Chunking)
- Preserves DOCX structure, headings, tables, and formatting

Requirements:
    pip install python-docx python-dotenv tqdm tenacity customtkinter requests openai

Usage:
    GUI Mode:   python translate_docx.py
    CLI Mode:   python translate_docx.py --input proposal_thai.docx --output proposal_english.tex --direction thai-to-english
"""

import argparse
import json
import logging
import os
import queue
import re
import sys
import threading
import time
import random
from pathlib import Path
from typing import Any, Callable, Dict, List, Tuple, Optional
from abc import ABC, abstractmethod
from datetime import datetime, timedelta

import customtkinter as ctk
import docx
import requests
from docx.document import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from dotenv import load_dotenv
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type
from tkinter import filedialog, messagebox
from tqdm import tqdm

# Try to import OpenAI for Typhoon support
try:
    from openai import OpenAI, RateLimitError
    TYPHOON_AVAILABLE = True
except ImportError:
    TYPHOON_AVAILABLE = False

# --- Constants & Prompts ---

# Translation prompts with language direction
TRANSLATION_PROMPTS = {
    "thai-to-english": """Translate the following Thai text into accurate, natural English.
Preserve meaning, technical terminology, abbreviations, equations, and units.
Return ONLY the translation, no explanations or markdown.""",

    "english-to-thai": """Translate the following English text into accurate, natural Thai.
Preserve meaning, technical terminology, abbreviations, equations, and units.
Return ONLY the translation, no explanations or markdown."""
}

# Language direction prefixes for user prompts
LANGUAGE_PREFIXES = {
    "thai-to-english": "Thai to English, ",
    "english-to-thai": "English to Thai, "
}

# Typhoon system prompt (exactly as used in playground)
TYPHOON_SYSTEM_PROMPT = """You are an AI assistant named Typhoon created by SCB 10X to be helpful, harmless, and honest. Typhoon is happy to help with analysis, question answering, math, coding, creative writing, teaching, role-play, general discussion, and all sorts of other tasks. Typhoon responds directly to all human messages without unnecessary affirmations or filler phrases like "Certainly!", "Of course!", "Absolutely!", "Great!", "Sure!", etc. Specifically, Typhoon avoids starting responses with the word "Certainly" in any way. Typhoon follows this information in all languages, and always responds to the user in the language they use or request. Typhoon is now being connected with a human. Write in fluid, conversational prose, Show genuine interest in understanding requests, Express appropriate emotions and empathy."""

# --- Helper Functions ---

# API Key persistence
API_KEY_FILE = "typhoon_api_key.txt"


def save_api_key(api_key: str):
    """Save API key to a text file."""
    try:
        with open(API_KEY_FILE, 'w') as f:
            f.write(api_key)
        return True
    except Exception as e:
        logger.error(f"Failed to save API key: {e}")
        return False


def load_api_key() -> Optional[str]:
    """Load API key from text file."""
    try:
        if os.path.exists(API_KEY_FILE):
            with open(API_KEY_FILE, 'r') as f:
                return f.read().strip()
    except Exception as e:
        logger.error(f"Failed to load API key: {e}")
    return None


def is_mostly_english(text: str) -> bool:
    """Determine if a text block is predominantly English (to skip translation)."""
    if not text or not text.strip():
        return True

    # Count Thai characters
    thai_chars = len(re.findall(r'[\u0E00-\u0E7F]', text))

    # Count English letters (a-z, A-Z)
    english_chars = len(re.findall(r'[a-zA-Z]', text))

    # Count total meaningful characters (Thai + English)
    total_meaningful = thai_chars + english_chars

    if total_meaningful == 0:
        return True  # Pure numbers/symbols

    # If more than 80% English characters, treat as English
    if english_chars / total_meaningful > 0.8:
        return True

    # If more than 15% Thai characters, treat as Thai
    if thai_chars / total_meaningful > 0.15:
        return False

    return english_chars / total_meaningful > 0.5


def detect_language(text: str) -> str:
    """Detect if text is mostly Thai or English."""
    if is_mostly_english(text):
        return "english"
    else:
        return "thai"


def estimate_processing_time(total_chunks: int, api_type: str, model: str) -> Tuple[int, str]:
    """Estimate processing time based on chunks and API type."""
    # Average time per chunk in seconds (estimated)
    avg_time_per_chunk = {
        "ollama": 5.0,      # Local model
        "typhoon": 3.0,     # Online model with rate limiting
        "openrouter": 4.0   # Online model
    }

    avg_time = avg_time_per_chunk.get(api_type.lower(), 4.0)

    # Adjust for model complexity
    if "30b" in model or "large" in model:
        avg_time *= 1.5
    elif "7b" in model or "small" in model:
        avg_time *= 0.7

    total_seconds = total_chunks * avg_time
    minutes = int(total_seconds // 60)
    seconds = int(total_seconds % 60)

    if minutes > 60:
        hours = minutes // 60
        minutes = minutes % 60
        return total_seconds, f"{hours}h {minutes}m"
    else:
        return total_seconds, f"{minutes}m {seconds}s"


def chunk_text(text: str, max_chunk_size: int = 3000) -> List[str]:
    """Split text into manageable chunks for the LLM, preferring sentence boundaries."""
    if len(text) <= max_chunk_size:
        return [text]

    chunks = []
    current_chunk = []
    current_length = 0

    # Split by paragraphs first (double newlines)
    paragraphs = re.split(r'\n\s*\n', text)

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        # If this paragraph alone is too long, split by sentences
        if len(para) > max_chunk_size:
            if current_chunk:
                chunks.append("\n\n".join(current_chunk))
                current_chunk = []
                current_length = 0

            # Split long paragraph by sentences
            sentences = re.split(r'(?<=[.!?])\s+', para)
            for sentence in sentences:
                if len(sentence) > max_chunk_size:
                    # If sentence is still too long, split by characters
                    if current_chunk:
                        chunks.append("\n\n".join(current_chunk))
                        current_chunk = []
                        current_length = 0
                    for i in range(0, len(sentence), max_chunk_size):
                        chunks.append(sentence[i:i+max_chunk_size])
                elif current_length + len(sentence) > max_chunk_size:
                    chunks.append("\n\n".join(current_chunk))
                    current_chunk = [sentence]
                    current_length = len(sentence)
                else:
                    current_chunk.append(sentence)
                    current_length += len(sentence) + 1
            continue

        # Check if adding this paragraph would exceed limit
        if current_length + len(para) > max_chunk_size and current_chunk:
            chunks.append("\n\n".join(current_chunk))
            current_chunk = [para]
            current_length = len(para)
        else:
            current_chunk.append(para)
            current_length += len(para) + 2  # +2 for the paragraph separator

    if current_chunk:
        chunks.append("\n\n".join(current_chunk))

    return chunks


def get_optimal_chunk_size(texts: List[str], target_chunks: int = None) -> int:
    """Calculate optimal chunk size based on text distribution."""
    if not texts:
        return 3000

    # Get lengths of all texts
    lengths = [len(t) for t in texts if t]
    if not lengths:
        return 3000

    # Calculate statistics
    avg_len = sum(lengths) / len(lengths)
    max_len = max(lengths)
    min_len = min(lengths)

    # If we have a target number of chunks, calculate chunk size
    if target_chunks and len(texts) > target_chunks:
        total_len = sum(lengths)
        chunk_size = max(500, int(total_len / target_chunks))
        # Round to nearest 500
        chunk_size = ((chunk_size + 250) // 500) * 500
        return min(chunk_size, 8000)  # Cap at 8000 characters

    # Smart chunk sizing based on text distribution
    if avg_len < 100:
        # Many small paragraphs - batch them together
        return 4000
    elif avg_len < 500:
        # Medium paragraphs
        return 3000
    elif avg_len < 1000:
        # Large paragraphs
        return 2000
    else:
        # Very large paragraphs
        return 1500


def replace_paragraph_text_preserve_format(paragraph: Paragraph, new_text: str) -> None:
    """Replace paragraph text while attempting to preserve the first run's formatting."""
    if not paragraph.runs:
        paragraph.text = new_text
        return

    for run in paragraph.runs[1:]:
        run.text = ""

    paragraph.runs[0].text = new_text


def clean_translation_response(text: str) -> str:
    """Clean up translation responses to remove extra text and keep the good format."""
    # Remove common introductory phrases
    patterns = [
        r'^Certainly[,\s]+',
        r'^Of course[,\s]+',
        r'^Absolutely[,\s]+',
        r'^Great[,\s]+',
        r'^Sure[,\s]+',
        r'^Here is[,\s]+',
        r'^Here\'s[,\s]+',
        r'^I will[,\s]+',
        r'^I can[,\s]+',
        r'^Let me[,\s]+',
        r'^Please find[,\s]+',
        r'^The translation is[:\s]+',
        r'^Translation:[,\s]+',
        r'^English translation:[:\s]+',
        r'^Here is the translation:[:\s]+',
        r'^I have translated[:\s]+',
        r'^I\'ve translated[:\s]+',
        r'^Here\'s the translation[:\s]+',
        r'^As requested,[,\s]+',
        r'^Sure thing,[,\s]+',
        r'^Thai to English,[,\s]+',
        r'^English to Thai,[,\s]+',
    ]

    for pattern in patterns:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE)

    # Remove common closing phrases
    closing_patterns = [
        r'\s*Let me know if you need any changes\.$',
        r'\s*Let me know if you need any further assistance\.$',
        r'\s*Please let me know if you need any revisions\.$',
        r'\s*I hope this helps\.$',
        r'\s*I hope this is helpful\.$',
        r'\s*Thank you for your request\.$',
        r'\s*Please let me know if you have any questions\.$',
        r'\s*Feel free to ask if you need anything else\.$',
        r'\s*Please let me know if you need further assistance\.$',
        r'\s*I look forward to your feedback\.$',
    ]

    for pattern in closing_patterns:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE)

    # Remove any markdown code blocks
    text = re.sub(r'^```[a-zA-Z]*\n?', '', text)
    text = re.sub(r'\n?```$', '', text)

    # Remove any lines that contain only Thai characters
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        # Skip lines that are pure Thai
        if re.match(r'^[\u0E00-\u0E7F\s]+$', line):
            continue
        cleaned_lines.append(line)

    text = '\n'.join(cleaned_lines)

    # Clean up extra whitespace
    text = re.sub(r'\n\s*\n', '\n\n', text)
    text = re.sub(r' +', ' ', text)

    return text.strip()


# --- LaTeX Export Functions ---

def escape_latex(text: str) -> str:
    """Escape special LaTeX characters."""
    latex_special = {
        '&': r'\&',
        '%': r'\%',
        '$': r'\$',
        '#': r'\#',
        '_': r'\_',
        '{': r'\{',
        '}': r'\}',
        '~': r'\textasciitilde{}',
        '^': r'\textasciicircum{}',
        '\\': r'\textbackslash{}',
        '<': r'\textless{}',
        '>': r'\textgreater{}',
    }
    # Escape backslash first
    text = text.replace('\\', '\\\\')
    for char, escaped in latex_special.items():
        text = text.replace(char, escaped)
    return text


def detect_heading_level(paragraph: Paragraph) -> Optional[int]:
    """Detect heading level from paragraph style."""
    style_name = paragraph.style.name.lower() if paragraph.style else ""

    # Common heading patterns
    heading_patterns = [
        (r'heading\s*1', 1),
        (r'heading\s*2', 2),
        (r'heading\s*3', 3),
        (r'heading\s*4', 4),
        (r'heading\s*5', 5),
        (r'title', 0),
        (r'subtitle', 0),
    ]

    for pattern, level in heading_patterns:
        if re.search(pattern, style_name):
            return level

    # Check if text looks like a heading (short, no ending punctuation)
    text = paragraph.text.strip()
    if len(text) < 100 and not re.search(r'[.!?]$', text):
        # Check if it's all caps or has numbers like "1. Introduction"
        if text.isupper() or re.match(r'^\d+\.', text):
            return 2

    return None


def detect_bold(paragraph: Paragraph) -> bool:
    """Check if paragraph is bold."""
    for run in paragraph.runs:
        if run.bold:
            return True
    return False


def paragraph_to_latex(paragraph: Paragraph, include_heading: bool = True) -> str:
    """Convert a paragraph to LaTeX with formatting preservation."""
    text = paragraph.text.strip()
    if not text:
        return ""

    # Escape LaTeX special characters
    text = escape_latex(text)

    # Check for heading
    if include_heading:
        heading_level = detect_heading_level(paragraph)
        if heading_level is not None:
            if heading_level == 0:
                return f"\\title{{{text}}}\n\\maketitle"
            else:
                return f"\\section{{{text}}}" if heading_level == 1 else f"\\subsection{{{text}}}"

    # Check for bold
    if detect_bold(paragraph):
        return f"\\textbf{{{text}}}"

    # Check for lists (bullet points)
    if text.startswith('•') or text.startswith('-') or text.startswith('*'):
        return f"\\begin{{itemize}}\n  \\item {text[1:].strip()}\n\\end{{itemize}}"

    # Regular paragraph
    return text


def docx_to_latex(doc: Document) -> str:
    """Convert a DOCX document to LaTeX format."""
    latex_parts = []

    # Document preamble
    latex_parts.append(r"\documentclass[12pt,a4paper]{article}")
    latex_parts.append(r"\usepackage[utf8]{inputenc}")
    latex_parts.append(r"\usepackage[T1]{fontenc}")
    latex_parts.append(r"\usepackage{geometry}")
    latex_parts.append(
        r"\geometry{left=2.5cm,right=2.5cm,top=2.5cm,bottom=2.5cm}")
    latex_parts.append(r"\usepackage{graphicx}")
    latex_parts.append(r"\usepackage{amsmath}")
    latex_parts.append(r"\usepackage{amssymb}")
    latex_parts.append(r"\usepackage{array}")
    latex_parts.append(r"\usepackage{booktabs}")
    latex_parts.append(r"\usepackage{hyperref}")
    latex_parts.append(
        r"\hypersetup{colorlinks=true, linkcolor=blue, urlcolor=blue}")
    latex_parts.append(r"\usepackage{parskip}")
    latex_parts.append(r"\setlength{\parindent}{0pt}")
    latex_parts.append(r"\setlength{\parskip}{1em}")
    latex_parts.append(r"\begin{document}")

    # Process paragraphs
    for para in doc.paragraphs:
        latex_text = paragraph_to_latex(para)
        if latex_text:
            latex_parts.append(latex_text)
            latex_parts.append("")  # Add blank line between paragraphs

    # Process tables
    for table in doc.tables:
        latex_parts.append(table_to_latex(table))

    # End document
    latex_parts.append(r"\end{document}")

    return "\n".join(latex_parts)


def table_to_latex(table: Table) -> str:
    """Convert a DOCX table to LaTeX format."""
    if not table.rows:
        return ""

    # Determine column count
    col_count = len(table.rows[0].cells)

    latex_parts = []
    latex_parts.append(r"\begin{table}[htbp]")
    latex_parts.append(r"\centering")
    latex_parts.append(
        r"\begin{tabular}{" + "|" + "|".join(["c"] * col_count) + "|}")
    latex_parts.append(r"\hline")

    for row_idx, row in enumerate(table.rows):
        row_cells = []
        for cell in row.cells:
            cell_text = escape_latex(cell.text.strip())
            row_cells.append(cell_text)
        latex_parts.append(" & ".join(row_cells) + r" \\")
        latex_parts.append(r"\hline")

    latex_parts.append(r"\end{tabular}")
    latex_parts.append(r"\end{table}")

    return "\n".join(latex_parts)


def save_as_latex(doc: Document, output_path: str) -> None:
    """Save document as LaTeX file."""
    latex_content = docx_to_latex(doc)
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(latex_content)


# --- Rate Limiter for Typhoon ---

class RateLimiter:
    """Rate limiter for API requests with token bucket algorithm."""

    def __init__(self, requests_per_second: int = 5, requests_per_minute: int = 200):
        self.requests_per_second = requests_per_second
        self.requests_per_minute = requests_per_minute
        self.min_interval = 1.0 / requests_per_second if requests_per_second > 0 else 0.2
        self.min_interval_per_minute = 60.0 / \
            requests_per_minute if requests_per_minute > 0 else 0.3
        self.interval = max(self.min_interval, self.min_interval_per_minute)
        self.last_request_time = 0
        self.lock = threading.Lock()
        self.request_times = []
        self.window_size = 60  # 60 second window for per-minute rate limiting

    def wait_if_needed(self):
        """Wait if necessary to respect rate limits."""
        with self.lock:
            current_time = time.time()

            # Check per-second rate limit
            time_since_last = current_time - self.last_request_time
            if time_since_last < self.interval:
                sleep_time = self.interval - time_since_last + 0.05  # Add small buffer
                if sleep_time > 0:
                    time.sleep(sleep_time)

            # Check per-minute rate limit
            now = time.time()
            # Remove requests older than 60 seconds
            self.request_times = [
                t for t in self.request_times if now - t < 60]

            if len(self.request_times) >= self.requests_per_minute:
                # Wait until the oldest request is 60 seconds old
                oldest = self.request_times[0]
                wait_time = 60 - (now - oldest) + 0.1
                if wait_time > 0:
                    time.sleep(wait_time)
                    # Clean up again after waiting
                    now = time.time()
                    self.request_times = [
                        t for t in self.request_times if now - t < 60]

            # Record this request
            self.request_times.append(time.time())
            self.last_request_time = time.time()


# --- Abstract Base API Client ---

logger = logging.getLogger("docx_translator")


class BaseLLMClient(ABC):
    """Abstract base class for LLM API clients."""

    @abstractmethod
    def chat_completion(self, model: str, system_prompt: str, user_text: str, stats: Dict[str, int]) -> str:
        """Make a chat completion request to the API."""
        pass


# --- OpenRouter API Client ---

class OpenRouterClient(BaseLLMClient):
    """OpenRouter API client using requests."""

    def __init__(self, api_key: str, site_url: str = "", site_title: str = ""):
        self.api_key = api_key
        self.site_url = site_url
        self.site_title = site_title
        self.base_url = "https://openrouter.ai/api/v1/chat/completions"

    def _get_headers(self) -> Dict[str, str]:
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json",
        }
        if self.site_url:
            headers["HTTP-Referer"] = self.site_url
        if self.site_title:
            headers["X-OpenRouter-Title"] = self.site_title
        return headers

    @retry(
        stop=stop_after_attempt(5),
        wait=wait_exponential(multiplier=1, min=2, max=10),
        retry=retry_if_exception_type(Exception),
        before_sleep=lambda retry_state: logger.warning(
            f"Retrying OpenRouter API call... Attempt {retry_state.attempt_number}"),
        reraise=True
    )
    def chat_completion(self, model: str, system_prompt: str, user_text: str, stats: Dict[str, int]) -> str:
        """Make a robust, retriable API call to OpenRouter."""
        stats["api_calls"] = stats.get("api_calls", 0) + 1

        model = model.strip()

        payload = {
            "model": model,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_text}
            ],
            "temperature": 0.1,
            "max_tokens": 8192,
        }

        try:
            logger.debug(f"Sending request to OpenRouter with model: {model}")
            response = requests.post(
                url=self.base_url,
                headers=self._get_headers(),
                data=json.dumps(payload),
                timeout=120
            )
            response.raise_for_status()

            result = response.json()

            if "error" in result:
                error_msg = result["error"].get(
                    "message", str(result["error"]))
                logger.error(f"OpenRouter API error: {error_msg}")
                stats["failures"] = stats.get("failures", 0) + 1
                raise Exception(f"OpenRouter API error: {error_msg}")

            content = result["choices"][0]["message"]["content"].strip()
            content = clean_translation_response(content)

            if content.startswith("```"):
                content = re.sub(r"^```[a-zA-Z]*\n?", "", content)
                content = re.sub(r"\n?```$", "", content)

            return content.strip()

        except requests.exceptions.RequestException as e:
            stats["failures"] = stats.get("failures", 0) + 1
            error_msg = str(e)

            if "400" in error_msg:
                logger.error(
                    f"Bad request - model '{model}' may not exist or is not accessible")
            elif "401" in error_msg:
                logger.error(
                    "Authentication failed - please check your API key")
            elif "429" in error_msg:
                logger.error("Rate limit exceeded - please wait and try again")
            else:
                logger.error(f"OpenRouter API request failed: {e}")

            if hasattr(e, 'response') and e.response:
                try:
                    error_data = e.response.json()
                    if "error" in error_data:
                        logger.error(f"API Error: {error_data['error']}")
                    else:
                        logger.error(f"Response: {error_data}")
                except:
                    logger.error(f"Response: {e.response.text[:200]}")
            raise


# --- Ollama API Client ---

class OllamaClient(BaseLLMClient):
    """Ollama local API client using requests."""

    def __init__(self, base_url: str = "http://localhost:11434"):
        self.base_url = base_url.rstrip('/')
        self.api_url = f"{self.base_url}/api/chat"

    def _test_connection(self) -> bool:
        try:
            response = requests.get(f"{self.base_url}/api/tags", timeout=5)
            return response.status_code == 200
        except:
            return False

    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=1, max=5),
        retry=retry_if_exception_type(Exception),
        before_sleep=lambda retry_state: logger.warning(
            f"Retrying Ollama API call... Attempt {retry_state.attempt_number}"),
        reraise=True
    )
    def chat_completion(self, model: str, system_prompt: str, user_text: str, stats: Dict[str, int]) -> str:
        stats["api_calls"] = stats.get("api_calls", 0) + 1

        if not self._test_connection():
            raise ConnectionError(
                f"Cannot connect to Ollama at {self.base_url}. "
                "Please ensure Ollama is running."
            )

        payload = {
            "model": model.strip(),
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_text}
            ],
            "stream": False,
            "options": {
                "temperature": 0.1,
                "num_predict": 8192,
            }
        }

        try:
            response = requests.post(
                url=self.api_url,
                json=payload,
                timeout=120
            )
            response.raise_for_status()

            result = response.json()
            content = result["message"]["content"].strip()
            content = clean_translation_response(content)

            if content.startswith("```"):
                content = re.sub(r"^```[a-zA-Z]*\n?", "", content)
                content = re.sub(r"\n?```$", "", content)

            return content.strip()

        except requests.exceptions.RequestException as e:
            stats["failures"] = stats.get("failures", 0) + 1
            if "404" in str(e):
                logger.error(
                    f"Model '{model}' not found in Ollama. Please pull it first:")
                logger.error(f"  ollama pull {model}")
            else:
                logger.error(f"Ollama API request failed: {e}")
            if hasattr(e, 'response') and e.response:
                logger.error(f"Response: {e.response.text[:200]}")
            raise
        except KeyError as e:
            stats["failures"] = stats.get("failures", 0) + 1
            logger.error(f"Unexpected Ollama response format: {e}")
            if 'response' in locals():
                logger.error(f"Response: {response.text[:200]}")
            raise


# --- Typhoon API Client (Online - using OpenAI SDK with Rate Limiting) ---

class TyphoonClient(BaseLLMClient):
    """Typhoon API client using OpenAI SDK with rate limiting."""

    def __init__(self, api_key: str):
        if not TYPHOON_AVAILABLE:
            raise ImportError(
                "OpenAI package is required for Typhoon. Install with: pip install openai")

        self.api_key = api_key
        self.client = OpenAI(
            api_key=api_key,
            base_url="https://api.opentyphoon.ai/v1",
            timeout=120.0,
            max_retries=0  # We'll handle retries ourselves
        )
        # Initialize rate limiter with default limits
        self.rate_limiter = RateLimiter(
            requests_per_second=5,
            requests_per_minute=200
        )
        self.last_rate_limit_time = 0
        self.rate_limit_retry_count = 0

    def _call_with_exponential_backoff(self, model: str, messages: List[Dict], max_retries: int = 5) -> str:
        """Call Typhoon API with exponential backoff for rate limits."""
        retries = 0

        while retries < max_retries:
            try:
                # Apply rate limiting before making the request
                self.rate_limiter.wait_if_needed()

                # Make the API call
                stream = self.client.chat.completions.create(
                    model=model,
                    messages=messages,
                    temperature=0.6,
                    max_completion_tokens=512,
                    top_p=0.6,
                    frequency_penalty=0,
                    stream=True
                )

                # Process the streaming response
                content_parts = []
                for chunk in stream:
                    if chunk.choices[0].delta.content is not None:
                        content_parts.append(chunk.choices[0].delta.content)

                content = ''.join(content_parts).strip()

                # Reset retry count on success
                self.rate_limit_retry_count = 0
                return content

            except RateLimitError as e:
                retries += 1
                self.rate_limit_retry_count = retries

                # Calculate backoff time with exponential backoff + jitter
                backoff_time = (2 ** retries) + random.random()

                logger.warning(
                    f"Rate limit exceeded (attempt {retries}/{max_retries}). "
                    f"Retrying in {backoff_time:.2f} seconds..."
                )

                if retries >= max_retries:
                    logger.error("Max retries exceeded for rate limit")
                    raise e

                time.sleep(backoff_time)

            except Exception as e:
                # For non-rate-limit errors, retry with shorter backoff
                retries += 1
                error_msg = str(e)

                if "authentication" in error_msg.lower() or "api_key" in error_msg.lower():
                    logger.error(
                        f"Authentication failed - please check your Typhoon API key")
                    raise e
                elif "model" in error_msg.lower() and "not found" in error_msg.lower():
                    logger.error(
                        f"Model '{model}' not found. Try 'typhoon-v2.5-30b-a3b-instruct'")
                    raise e
                else:
                    logger.warning(
                        f"API error (attempt {retries}/{max_retries}): {e}")

                    if retries >= max_retries:
                        raise e

                    # Exponential backoff for other errors too
                    backoff_time = (1.5 ** retries) + random.random() * 0.5
                    time.sleep(min(backoff_time, 10))

        raise Exception("Max retries exceeded")

    def chat_completion(self, model: str, system_prompt: str, user_text: str, stats: Dict[str, int]) -> str:
        """Make a robust, retriable API call to Typhoon with rate limiting."""
        stats["api_calls"] = stats.get("api_calls", 0) + 1

        model = model.strip()

        # Use the system prompt and user prompt exactly like the playground
        messages = [
            {"role": "system", "content": TYPHOON_SYSTEM_PROMPT},
            {"role": "user", "content": user_text}
        ]

        try:
            logger.debug(f"Sending request to Typhoon with model: {model}")

            content = self._call_with_exponential_backoff(model, messages)

            # Clean up the response
            content = clean_translation_response(content)

            if content.startswith("```"):
                content = re.sub(r"^```[a-zA-Z]*\n?", "", content)
                content = re.sub(r"\n?```$", "", content)

            return content.strip()

        except Exception as e:
            stats["failures"] = stats.get("failures", 0) + 1
            logger.error(f"Typhoon API request failed after retries: {e}")
            raise


# --- API Client Factory ---

def create_llm_client(api_type: str, **kwargs) -> BaseLLMClient:
    if api_type.lower() == "openrouter":
        api_key = kwargs.get("api_key")
        if not api_key:
            raise ValueError("API key required for OpenRouter")
        return OpenRouterClient(
            api_key=api_key,
            site_url=kwargs.get("site_url", "https://github.com"),
            site_title=kwargs.get("site_title", "DOCX Translator")
        )
    elif api_type.lower() == "ollama":
        base_url = kwargs.get("base_url", "http://localhost:11434")
        return OllamaClient(base_url=base_url)
    elif api_type.lower() == "typhoon":
        api_key = kwargs.get("api_key")
        if not api_key:
            raise ValueError("API key required for Typhoon")
        return TyphoonClient(api_key=api_key)
    else:
        raise ValueError(f"Unsupported API type: {api_type}")


# --- Translation Pipeline (Single Stage) ---

def translate_text(client: BaseLLMClient, text: str, model: str, direction: str, stats: Dict[str, int]) -> str:
    """Execute single-stage translation with language direction."""
    chunks = chunk_text(text, max_chunk_size=3000)

    # Get the appropriate prompt and prefix
    system_prompt = TRANSLATION_PROMPTS.get(
        direction, TRANSLATION_PROMPTS["thai-to-english"])
    language_prefix = LANGUAGE_PREFIXES.get(
        direction, LANGUAGE_PREFIXES["thai-to-english"])

    results = []
    total_chunks = len(chunks)

    for i, chunk in enumerate(chunks):
        logger.debug(
            f"Translation chunk {i+1}/{total_chunks}: {len(chunk)} chars")
        # Add language direction prefix to the user text
        user_text_with_prefix = f"{language_prefix}{chunk}"
        res = client.chat_completion(
            model, system_prompt, user_text_with_prefix, stats)
        results.append(res)

        # Add a small delay between chunks to help with rate limiting
        if i < total_chunks - 1:
            time.sleep(0.3)

    return " ".join(results)


class CallbackLogHandler(logging.Handler):
    def __init__(self, callback: Callable[[str], None]):
        super().__init__()
        self.callback = callback

    def emit(self, record: logging.LogRecord):
        msg = self.format(record)
        self.callback(msg)


class TranslationProgress:
    """Track translation progress with time estimation."""

    def __init__(self, total_items: int):
        self.total_items = total_items
        self.processed_items = 0
        self.start_time = None
        self.last_update_time = None
        self.estimated_time_per_item = None
        self.eta_seconds = 0
        self.history = []

    def start(self):
        self.start_time = time.time()
        self.last_update_time = self.start_time

    def update(self, processed: int):
        self.processed_items = processed
        current_time = time.time()

        if self.processed_items > 0 and self.processed_items <= self.total_items:
            elapsed = current_time - self.start_time
            self.estimated_time_per_item = elapsed / self.processed_items
            remaining_items = self.total_items - self.processed_items
            self.eta_seconds = remaining_items * self.estimated_time_per_item

        self.last_update_time = current_time

    def get_eta(self) -> str:
        """Get formatted ETA string."""
        if self.eta_seconds <= 0:
            return "Calculating..."

        if self.eta_seconds < 60:
            return f"{int(self.eta_seconds)}s"
        elif self.eta_seconds < 3600:
            minutes = int(self.eta_seconds // 60)
            seconds = int(self.eta_seconds % 60)
            return f"{minutes}m {seconds}s"
        else:
            hours = int(self.eta_seconds // 3600)
            minutes = int((self.eta_seconds % 3600) // 60)
            return f"{hours}h {minutes}m"

    def get_percentage(self) -> float:
        if self.total_items <= 0:
            return 0
        return (self.processed_items / self.total_items) * 100

    def get_progress_text(self) -> str:
        if self.processed_items == 0:
            return "Starting..."

        pct = self.get_percentage()
        eta = self.get_eta()
        return f"{self.processed_items}/{self.total_items} ({pct:.1f}%) - ETA: {eta}"


def run_translation_task(
    input_path: str,
    output_path: str,
    model: str,
    api_type: str,
    direction: str,
    api_key: Optional[str] = None,
    ollama_url: str = "http://localhost:11434",
    output_format: str = "docx",
    progress_callback: Callable[[int, int, str], None] = None,
    log_callback: Callable[[str], None] = None,
    stop_event: threading.Event = None
) -> Dict[str, int]:

    logger.handlers = []
    logger.setLevel(logging.INFO)

    fh = logging.FileHandler("translation.log", encoding="utf-8")
    fh.setFormatter(logging.Formatter(
        '%(asctime)s [%(levelname)s] %(message)s'))
    logger.addHandler(fh)

    if log_callback:
        ch = CallbackLogHandler(log_callback)
        ch.setFormatter(logging.Formatter('[%(levelname)s] %(message)s'))
        logger.addHandler(ch)

    model = model.strip()

    logger.info(f"Initializing {api_type} client...")
    logger.info(f"Translation direction: {direction}")
    logger.info(f"Output format: {output_format}")
    try:
        if api_type.lower() == "openrouter":
            if not api_key:
                raise ValueError("OpenRouter requires an API key")
            client = create_llm_client(
                api_type="openrouter",
                api_key=api_key,
                site_url="https://github.com",
                site_title="DOCX Translator"
            )
        elif api_type.lower() == "typhoon":
            if not api_key:
                raise ValueError("Typhoon requires an API key")
            client = create_llm_client(
                api_type="typhoon",
                api_key=api_key
            )
            logger.info(
                "Typhoon rate limiting enabled (5 req/sec, 200 req/min)")
        else:
            client = create_llm_client(
                api_type="ollama",
                base_url=ollama_url
            )
        logger.info(f"Using model: {model}")
    except Exception as e:
        logger.error(f"Failed to initialize client: {e}")
        raise

    logger.info(f"Loading document: {input_path}")
    doc = docx.Document(input_path)

    stats = {
        "paragraphs_translated": 0,
        "paragraphs_skipped": 0,
        "table_cells_translated": 0,
        "table_cells_skipped": 0,
        "api_calls": 0,
        "failures": 0,
        "total_segments": 0,
        "total_chunks": 0
    }

    to_process: List[Tuple[Paragraph, str]] = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue

        # For Thai->English, skip English text
        # For English->Thai, skip Thai text
        if direction == "thai-to-english":
            if is_mostly_english(text):
                stats["paragraphs_skipped"] += 1
                continue
        else:  # english-to-thai
            if not is_mostly_english(text):
                stats["paragraphs_skipped"] += 1
                continue

        to_process.append((p, "paragraph"))

    def get_table_paragraphs(tbl: Table) -> List[Tuple[Paragraph, str]]:
        paras = []
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = p.text.strip()
                    if not text:
                        continue
                    if direction == "thai-to-english":
                        if is_mostly_english(text):
                            stats["table_cells_skipped"] += 1
                            continue
                    else:  # english-to-thai
                        if not is_mostly_english(text):
                            stats["table_cells_skipped"] += 1
                            continue
                    paras.append((p, "table_cell"))
                for nested_table in cell.tables:
                    paras.extend(get_table_paragraphs(nested_table))
        return paras

    for table in doc.tables:
        to_process.extend(get_table_paragraphs(table))

    total_segments = len(to_process)
    stats["total_segments"] = total_segments

    if total_segments == 0:
        logger.info("No text found to translate for the selected direction")
        if progress_callback:
            progress_callback(1, 1, "Complete - No text to translate")
        return stats

    # Calculate optimal chunk size
    segment_texts = [p.text.strip() for p, _ in to_process]
    chunk_size = get_optimal_chunk_size(segment_texts, target_chunks=50)
    logger.info(f"Optimal chunk size: {chunk_size} characters")

    # Estimate total chunks
    total_chunks = 0
    for text in segment_texts:
        chunks = chunk_text(text, max_chunk_size=chunk_size)
        total_chunks += len(chunks)
    stats["total_chunks"] = total_chunks

    # Estimate time
    est_seconds, est_time_str = estimate_processing_time(
        total_chunks, api_type, model)
    logger.info(
        f"Estimated processing time: {est_time_str} ({total_chunks} chunks)")

    if progress_callback:
        progress_callback(
            0, total_chunks, f"Estimating... {est_time_str} expected")

    # Progress tracker
    progress = TranslationProgress(total_chunks)
    progress.start()

    logger.info(
        f"Found {total_segments} segments, {total_chunks} chunks requiring translation.")

    chunk_counter = 0
    for i, (p, p_type) in enumerate(to_process):
        if stop_event and stop_event.is_set():
            logger.warning("Translation aborted by user.")
            break

        text = p.text.strip()

        # Chunk the text
        chunks = chunk_text(text, max_chunk_size=chunk_size)

        try:
            # Translate each chunk
            translated_parts = []
            for chunk_idx, chunk in enumerate(chunks):
                if stop_event and stop_event.is_set():
                    break

                chunk_counter += 1
                progress.update(chunk_counter)

                # Get current ETA and progress
                eta_text = progress.get_progress_text()

                if progress_callback:
                    progress_callback(chunk_counter, total_chunks, eta_text)

                logger.info(
                    f"Segment {i+1}/{total_segments}, Chunk {chunk_idx+1}/{len(chunks)} "
                    f"({len(chunk)} chars) - {eta_text}"
                )

                # Translate this chunk
                system_prompt = TRANSLATION_PROMPTS.get(direction)
                language_prefix = LANGUAGE_PREFIXES.get(direction)
                user_text_with_prefix = f"{language_prefix}{chunk}"
                res = client.chat_completion(
                    model, system_prompt, user_text_with_prefix, stats)
                translated_parts.append(res)

                # Small delay between chunks in same segment
                if chunk_idx < len(chunks) - 1:
                    time.sleep(0.2)

            if stop_event and stop_event.is_set():
                break

            translated_text = " ".join(translated_parts)

            if translated_text and translated_text != text:
                replace_paragraph_text_preserve_format(p, translated_text)
                if p_type == "paragraph":
                    stats["paragraphs_translated"] += 1
                else:
                    stats["table_cells_translated"] += 1
            else:
                if not translated_text:
                    logger.warning(
                        f"Empty translation for segment {i+1} (text: '{text[:50]}...')")
                elif translated_text == text:
                    logger.warning(
                        f"⚠️ Unchanged translation for segment {i+1} (text: '{text[:50]}...') - already English?")
                else:
                    logger.warning(f"Translation issue for segment {i+1}")

        except Exception as e:
            logger.error(f"Failed to translate segment {i+1}: {e}")
            stats["failures"] += 1

    if not (stop_event and stop_event.is_set()):
        logger.info(f"Saving translated document to: {output_path}")

        # Save in the appropriate format
        if output_format.lower() == "latex" or output_path.endswith('.tex'):
            save_as_latex(doc, output_path)
            logger.info("Translation complete! LaTeX file saved.")
        else:
            doc.save(output_path)
            logger.info("Translation complete! DOCX file saved.")

        if progress_callback:
            progress_callback(total_chunks, total_chunks, "Complete! ✅")

    return stats


# --- Modern GUI (2026) ---

class TranslatorGUI(ctk.CTk):
    def __init__(self):
        super().__init__()

        self.log_queue = queue.Queue()
        self.progress_queue = queue.Queue()
        self.stats_queue = queue.Queue()
        self.stop_event = threading.Event()
        self.is_running = False
        self.worker_thread = None

        self.title("AI Thai-English Translator")
        self.geometry("1100x850")
        self.minsize(950, 750)

        ctk.set_appearance_mode("System")
        ctk.set_default_color_theme("blue")

        self.build_ui()
        self.poll_queues()

    def change_theme(self, theme: str):
        ctk.set_appearance_mode(theme)

    def build_ui(self):
        top_bar = ctk.CTkFrame(self, fg_color="transparent")
        top_bar.pack(fill="x", padx=20, pady=(15, 5))

        ctk.CTkLabel(
            top_bar,
            text="AI LLM Thai-English Translator",
            font=ctk.CTkFont(size=22, weight="bold")
        ).pack(side="left")

        self.theme_var = ctk.StringVar(value="System")
        theme_switch = ctk.CTkSegmentedButton(
            top_bar,
            values=["System", "Light", "Dark"],
            variable=self.theme_var,
            command=self.change_theme,
            width=200
        )
        theme_switch.pack(side="right")

        main_frame = ctk.CTkFrame(self, fg_color="transparent")
        main_frame.pack(fill="both", expand=True, padx=20, pady=10)

        left_panel = ctk.CTkFrame(main_frame, corner_radius=15)
        left_panel.pack(side="left", fill="both", expand=True, padx=(0, 10))

        right_panel = ctk.CTkFrame(main_frame, corner_radius=15)
        right_panel.pack(side="right", fill="both", expand=True, padx=(10, 0))

        ctk.CTkLabel(
            left_panel,
            text="Configuration",
            font=ctk.CTkFont(size=18, weight="bold")
        ).pack(pady=(20, 15), padx=20, anchor="w")

        # Input File
        self.input_var = ctk.StringVar()
        ctk.CTkLabel(left_panel, text="Input DOCX:").pack(padx=20, anchor="w")
        input_frame = ctk.CTkFrame(left_panel, fg_color="transparent")
        input_frame.pack(fill="x", padx=20, pady=(0, 15))
        ctk.CTkEntry(input_frame, textvariable=self.input_var).pack(
            side="left", fill="x", expand=True, padx=(0, 5)
        )
        ctk.CTkButton(input_frame, text="Browse", width=80,
                      command=self.browse_input).pack(side="right")

        # Output File
        self.output_var = ctk.StringVar(value="proposal_english.docx")
        ctk.CTkLabel(left_panel, text="Output File:").pack(padx=20, anchor="w")
        output_frame = ctk.CTkFrame(left_panel, fg_color="transparent")
        output_frame.pack(fill="x", padx=20, pady=(0, 15))
        ctk.CTkEntry(output_frame, textvariable=self.output_var).pack(
            side="left", fill="x", expand=True, padx=(0, 5)
        )
        ctk.CTkButton(output_frame, text="Browse", width=80,
                      command=self.browse_output).pack(side="right")

        # Output Format
        self.format_var = ctk.StringVar(value="DOCX")
        ctk.CTkLabel(left_panel, text="Output Format:").pack(
            padx=20, anchor="w")
        format_frame = ctk.CTkFrame(left_panel, fg_color="transparent")
        format_frame.pack(fill="x", padx=20, pady=(0, 10))

        format_options = ["DOCX", "LaTeX (.tex)"]
        self.format_menu = ctk.CTkSegmentedButton(
            format_frame,
            values=format_options,
            variable=self.format_var,
            command=self.on_format_change,
            width=300
        )
        self.format_menu.pack(side="left", fill="x", expand=True)

        # API Type
        self.api_type_var = ctk.StringVar(value="Typhoon")
        ctk.CTkLabel(left_panel, text="API Backend:").pack(padx=20, anchor="w")
        api_type_frame = ctk.CTkFrame(left_panel, fg_color="transparent")
        api_type_frame.pack(fill="x", padx=20, pady=(0, 10))

        api_options = ["Ollama", "Typhoon", "OpenRouter"]

        self.api_type_menu = ctk.CTkSegmentedButton(
            api_type_frame,
            values=api_options,
            variable=self.api_type_var,
            command=self.on_api_type_change,
            width=300
        )
        self.api_type_menu.pack(side="left", fill="x", expand=True)

        # Language Direction
        self.direction_var = ctk.StringVar(value="Thai to English")
        ctk.CTkLabel(left_panel, text="Translation Direction:").pack(
            padx=20, anchor="w")
        direction_frame = ctk.CTkFrame(left_panel, fg_color="transparent")
        direction_frame.pack(fill="x", padx=20, pady=(0, 10))

        direction_options = ["Thai to English", "English to Thai"]
        self.direction_menu = ctk.CTkSegmentedButton(
            direction_frame,
            values=direction_options,
            variable=self.direction_var,
            command=self.on_direction_change,
            width=300
        )
        self.direction_menu.pack(side="left", fill="x", expand=True)

        # Model
        self.model_var = ctk.StringVar(
            value="typhoon-v2.5-30b-a3b-instruct")
        ctk.CTkLabel(left_panel, text="Model Name:").pack(padx=20, anchor="w")
        self.model_entry = ctk.CTkEntry(
            left_panel, textvariable=self.model_var)
        self.model_entry.pack(fill="x", padx=20, pady=(0, 10))

        # API Key
        self.api_key_var = ctk.StringVar()
        self.api_key_frame = ctk.CTkFrame(left_panel, fg_color="transparent")
        self.api_key_frame.pack(fill="x", padx=20, pady=(0, 10))

        self.api_key_label = ctk.CTkLabel(self.api_key_frame, text="API Key:")
        self.api_key_label.pack(anchor="w")
        api_key_row = ctk.CTkFrame(self.api_key_frame, fg_color="transparent")
        api_key_row.pack(fill="x", pady=(5, 0))
        self.api_entry = ctk.CTkEntry(
            api_key_row, textvariable=self.api_key_var, show="*"
        )
        self.api_entry.pack(side="left", fill="x", expand=True, padx=(0, 5))

        api_key_btn_frame = ctk.CTkFrame(api_key_row, fg_color="transparent")
        api_key_btn_frame.pack(side="right")

        ctk.CTkButton(api_key_btn_frame, text="Load .env", width=80,
                      command=self.load_env).pack(side="left", padx=(0, 5))
        ctk.CTkButton(api_key_btn_frame, text="Save Key", width=80,
                      command=self.save_api_key).pack(side="left")

        # Ollama URL
        self.ollama_url_var = ctk.StringVar(value="http://localhost:11434")
        self.ollama_url_frame = ctk.CTkFrame(
            left_panel, fg_color="transparent")

        ctk.CTkLabel(self.ollama_url_frame,
                     text="Ollama Base URL:").pack(anchor="w")
        self.ollama_url_entry = ctk.CTkEntry(
            self.ollama_url_frame, textvariable=self.ollama_url_var
        )
        self.ollama_url_entry.pack(fill="x", pady=(5, 0))

        # Rate Limit Info (for Typhoon)
        self.rate_limit_label = ctk.CTkLabel(
            left_panel,
            text="",
            font=ctk.CTkFont(size=11),
            text_color="gray"
        )
        self.rate_limit_label.pack(padx=20, pady=(0, 10), anchor="w")

        self.on_api_type_change(self.api_type_var.get())

        # Action Buttons
        action_frame = ctk.CTkFrame(left_panel, fg_color="transparent")
        action_frame.pack(fill="x", padx=20, pady=(15, 15))

        self.start_btn = ctk.CTkButton(
            action_frame,
            text="Start Translation",
            fg_color="#2fa84f",
            hover_color="#258a40",
            command=self.start_translation,
            height=40
        )
        self.start_btn.pack(side="left", expand=True, fill="x", padx=(0, 5))

        self.stop_btn = ctk.CTkButton(
            action_frame,
            text="Stop",
            fg_color="#d93636",
            hover_color="#b32c2c",
            command=self.stop_translation,
            state="disabled",
            height=40
        )
        self.stop_btn.pack(side="right", expand=True, fill="x", padx=(5, 0))

        # Progress
        self.progress_label = ctk.CTkLabel(left_panel, text="Idle")
        self.progress_label.pack(padx=20, anchor="w", pady=(10, 0))
        self.progressbar = ctk.CTkProgressBar(left_panel)
        self.progressbar.pack(fill="x", padx=20, pady=(5, 5))
        self.progressbar.set(0)

        # ETA Label
        self.eta_label = ctk.CTkLabel(
            left_panel,
            text="",
            font=ctk.CTkFont(size=12),
            text_color="gray"
        )
        self.eta_label.pack(padx=20, anchor="w", pady=(0, 10))

        # Right Panel - Console
        ctk.CTkLabel(
            right_panel,
            text="Translation Console",
            font=ctk.CTkFont(size=18, weight="bold")
        ).pack(pady=(20, 10), padx=20, anchor="w")

        self.log_text = ctk.CTkTextbox(
            right_panel,
            fg_color="#1e1e1e",
            text_color="#ffffff",
            font=ctk.CTkFont(family="Consolas", size=12),
            wrap="word",
            corner_radius=10
        )
        self.log_text.pack(fill="both", expand=True, padx=20, pady=(0, 20))
        self.log_text.configure(state="disabled")

        # Load saved API key
        saved_key = load_api_key()
        if saved_key:
            self.api_key_var.set(saved_key)

    def on_format_change(self, format_type: str):
        """Handle format change - update output file extension."""
        input_path = self.input_var.get()
        if input_path:
            base = os.path.basename(input_path)
            name, ext = os.path.splitext(base)
            direction = self.direction_var.get()

            if format_type == "LaTeX (.tex)":
                suffix = "_english.tex" if direction == "Thai to English" else "_thai.tex"
                self.output_var.set(os.path.join(
                    os.path.dirname(input_path), f"{name}{suffix}"
                ))
            else:
                suffix = "_english.docx" if direction == "Thai to English" else "_thai.docx"
                self.output_var.set(os.path.join(
                    os.path.dirname(input_path), f"{name}{suffix}"
                ))

    def on_direction_change(self, direction: str):
        """Handle direction change - auto-suggest output filename."""
        input_path = self.input_var.get()
        if input_path:
            base = os.path.basename(input_path)
            name, ext = os.path.splitext(base)
            format_type = self.format_var.get()

            if format_type == "LaTeX (.tex)":
                suffix = "_english.tex" if direction == "Thai to English" else "_thai.tex"
            else:
                suffix = "_english.docx" if direction == "Thai to English" else "_thai.docx"

            self.output_var.set(os.path.join(
                os.path.dirname(input_path), f"{name}{suffix}"
            ))

    def on_api_type_change(self, api_type: str):
        if api_type.lower() == "openrouter":
            self.api_key_label.configure(text="OpenRouter API Key:")
            self.api_key_frame.pack(fill="x", padx=20, pady=(0, 10))
            self.ollama_url_frame.pack_forget()
            self.model_var.set("deepseek/deepseek-v4-flash")
            self.rate_limit_label.configure(text="")
        elif api_type.lower() == "typhoon":
            self.api_key_label.configure(text="Typhoon API Key:")
            self.api_key_frame.pack(fill="x", padx=20, pady=(0, 10))
            self.ollama_url_frame.pack_forget()
            self.model_var.set("typhoon-v2.5-30b-a3b-instruct")
            self.rate_limit_label.configure(
                text="Rate Limits: 5 req/sec | 200 req/min (with exponential backoff)"
            )
        else:  # Ollama
            self.api_key_frame.pack_forget()
            self.ollama_url_frame.pack(fill="x", padx=20, pady=(0, 10))
            self.model_var.set("scb10x/typhoon-translate1.5-4b:latest")
            self.rate_limit_label.configure(text="")

    def browse_input(self):
        filepath = filedialog.askopenfilename(
            filetypes=[("Word Documents", "*.docx")])
        if filepath:
            self.input_var.set(filepath)
            self.on_direction_change(self.direction_var.get())

    def browse_output(self):
        format_type = self.format_var.get()
        if format_type == "LaTeX (.tex)":
            filetypes = [("LaTeX Files", "*.tex")]
            defaultextension = ".tex"
        else:
            filetypes = [("Word Documents", "*.docx")]
            defaultextension = ".docx"

        filepath = filedialog.asksaveasfilename(
            defaultextension=defaultextension,
            filetypes=filetypes
        )
        if filepath:
            self.output_var.set(filepath)

    def load_env(self):
        load_dotenv()
        api_type = self.api_type_var.get().lower()

        if api_type == "openrouter":
            key = os.getenv("OPENROUTER_API_KEY", "")
        elif api_type == "typhoon":
            key = os.getenv("TYPHOON_API_KEY", "")
        else:
            key = ""

        if key:
            self.api_key_var.set(key)
            messagebox.showinfo(
                "Success", f"API Key loaded from .env for {api_type}")
        else:
            messagebox.showwarning(
                "Warning", f"No {api_type.upper()}_API_KEY found in .env"
            )

    def save_api_key(self):
        """Save current API key to file."""
        key = self.api_key_var.get()
        if not key:
            messagebox.showwarning("Warning", "No API key to save.")
            return

        if save_api_key(key):
            messagebox.showinfo("Success", "API key saved successfully!")
        else:
            messagebox.showerror("Error", "Failed to save API key.")

    def start_translation(self):
        input_path = self.input_var.get()
        output_path = self.output_var.get()
        model = self.model_var.get()
        api_type = self.api_type_var.get()
        direction = self.direction_var.get()
        format_type = self.format_var.get()

        direction_key = "thai-to-english" if direction == "Thai to English" else "english-to-thai"
        output_format = "latex" if format_type == "LaTeX (.tex)" else "docx"

        if not input_path or not os.path.exists(input_path):
            messagebox.showerror("Error", "Please select a valid input file.")
            return

        if not output_path:
            messagebox.showerror("Error", "Please specify an output file.")
            return

        api_key = None
        ollama_url = None

        if api_type.lower() in ["openrouter", "typhoon"]:
            api_key = self.api_key_var.get()
            if not api_key:
                messagebox.showerror(
                    "Error", f"Please provide an {api_type} API key.")
                return
        else:
            ollama_url = self.ollama_url_var.get()
            if not ollama_url:
                messagebox.showerror(
                    "Error", "Please provide an Ollama base URL.")
                return

        self.is_running = True
        self.stop_event.clear()
        self.start_btn.configure(state="disabled")
        self.stop_btn.configure(state="normal")
        self.progressbar.set(0)
        self.progress_label.configure(text="Starting...")
        self.eta_label.configure(text="")

        self.log_text.configure(state="normal")
        self.log_text.delete("1.0", ctk.END)
        self.log_text.configure(state="disabled")

        self.worker_thread = threading.Thread(
            target=self.run_worker,
            args=(input_path, output_path, model, api_type,
                  direction_key, output_format, api_key, ollama_url),
            daemon=True
        )
        self.worker_thread.start()

    def run_worker(self, input_path, output_path, model, api_type, direction, output_format, api_key, ollama_url):
        try:
            stats = run_translation_task(
                input_path=input_path,
                output_path=output_path,
                model=model,
                api_type=api_type,
                direction=direction,
                output_format=output_format,
                api_key=api_key,
                ollama_url=ollama_url,
                progress_callback=lambda c, t, status: self.progress_queue.put(
                    (c, t, status)),
                log_callback=lambda m: self.log_queue.put(m),
                stop_event=self.stop_event
            )
            self.stats_queue.put(stats)
        except Exception as e:
            self.log_queue.put(f"[CRITICAL] Worker failed: {e}")
            self.stats_queue.put({"error": str(e)})
        finally:
            self.is_running = False

    def stop_translation(self):
        if self.is_running:
            self.stop_event.set()
            self.progress_label.configure(text="Stopping...")
            self.stop_btn.configure(state="disabled")
            self.log_queue.put(
                "[INFO] Stop signal sent. Waiting for current operation to finish...")

    def poll_queues(self):
        # Process log messages
        try:
            while True:
                msg = self.log_queue.get_nowait()
                self.log_text.configure(state="normal")
                self.log_text.insert(ctk.END, msg + "\n")
                self.log_text.see(ctk.END)
                self.log_text.configure(state="disabled")
        except queue.Empty:
            pass

        # Process progress updates
        try:
            while True:
                current, total, status = self.progress_queue.get_nowait()
                if total > 0:
                    self.progressbar.set(current / total)
                    self.progress_label.configure(text=status)

                    # Extract ETA from status if present
                    if "ETA:" in status:
                        eta_part = status.split("ETA:")[-1].strip()
                        self.eta_label.configure(text=f"⏱️ ETA: {eta_part}")
                    elif "Complete" in status:
                        self.eta_label.configure(text="✅ Done!")
                    else:
                        self.eta_label.configure(text=f"⏳ {current}/{total}")
                else:
                    self.progressbar.set(0)
                    self.progress_label.configure(text=status)
        except queue.Empty:
            pass

        # Process stats
        try:
            stats = self.stats_queue.get_nowait()
            self.on_translation_complete(stats)
        except queue.Empty:
            pass

        self.after(100, self.poll_queues)

    def on_translation_complete(self, stats: Dict):
        self.is_running = False
        self.start_btn.configure(state="normal")
        self.stop_btn.configure(state="disabled")
        self.stop_event.clear()

        if "error" in stats:
            messagebox.showerror(
                "Error", f"Translation failed:\n{stats['error']}")
            self.progress_label.configure(text="Failed")
            self.eta_label.configure(text="❌ Failed")
        else:
            self.progressbar.set(1.0)
            self.progress_label.configure(text="Completed!")
            self.eta_label.configure(text="✅ Done!")
            self.show_stats_modal(stats)

    def show_stats_modal(self, stats: Dict):
        modal = ctk.CTkToplevel(self)
        modal.title("Translation Statistics")
        modal.geometry("450x450")
        modal.transient(self)
        modal.grab_set()

        ctk.CTkLabel(
            modal,
            text="✅ Translation Complete!",
            font=ctk.CTkFont(size=20, weight="bold")
        ).pack(pady=(20, 10))

        frame = ctk.CTkFrame(modal, fg_color="transparent")
        frame.pack(fill="both", expand=True, padx=20, pady=10)

        display_order = [
            ("total_segments", "Total Segments"),
            ("total_chunks", "Total Chunks"),
            ("paragraphs_translated", "Paragraphs Translated"),
            ("paragraphs_skipped", "Paragraphs Skipped"),
            ("table_cells_translated", "Table Cells Translated"),
            ("table_cells_skipped", "Table Cells Skipped"),
            ("api_calls", "API Calls"),
            ("failures", "Failures"),
        ]

        for key, label in display_order:
            if key in stats:
                row = ctk.CTkFrame(frame, fg_color="transparent")
                row.pack(fill="x", pady=3)
                ctk.CTkLabel(row, text=label, anchor="w").pack(
                    side="left", fill="x", expand=True)
                ctk.CTkLabel(row, text=str(stats[key]), font=ctk.CTkFont(
                    weight="bold")).pack(side="right")

        ctk.CTkButton(modal, text="Close", command=modal.destroy,
                      width=120).pack(pady=20)


# --- Main Entry Point ---

def main():
    parser = argparse.ArgumentParser(
        description="Translate DOCX using Ollama, OpenRouter, or Typhoon."
    )
    parser.add_argument("--input", help="Input DOCX file path.")
    parser.add_argument("--output", help="Output file path (.docx or .tex).")
    parser.add_argument(
        "--model",
        default="typhoon-v2.5-30b-a3b-instruct",
        help="Model name."
    )
    parser.add_argument(
        "--api-type",
        choices=["openrouter", "ollama", "typhoon"],
        default="typhoon",
        help="API backend to use."
    )
    parser.add_argument(
        "--direction",
        choices=["thai-to-english", "english-to-thai"],
        default="thai-to-english",
        help="Translation direction."
    )
    parser.add_argument(
        "--format",
        choices=["docx", "latex"],
        default="docx",
        help="Output format (docx or latex)."
    )
    parser.add_argument(
        "--ollama-url",
        default="http://localhost:11434",
        help="Ollama base URL (default: http://localhost:11434)."
    )
    parser.add_argument("--gui", action="store_true", help="Force launch GUI.")
    parser.add_argument("--verbose", action="store_true",
                        help="Enable verbose logging.")

    args = parser.parse_args()

    if args.gui or not args.input or not args.output:
        app = TranslatorGUI()
        app.mainloop()
    else:
        load_dotenv()

        api_key = None
        ollama_url = None

        if args.api_type == "openrouter":
            api_key = os.getenv("OPENROUTER_API_KEY")
            if not api_key:
                api_key = load_api_key()
            if not api_key:
                print("Error: OPENROUTER_API_KEY not found in environment or .env file.")
                sys.exit(1)
        elif args.api_type == "typhoon":
            api_key = os.getenv("TYPHOON_API_KEY")
            if not api_key:
                api_key = load_api_key()
            if not api_key:
                print("Error: TYPHOON_API_KEY not found in environment or .env file.")
                sys.exit(1)
        else:
            ollama_url = args.ollama_url

        if args.verbose:
            logger.setLevel(logging.DEBUG)
        else:
            logger.setLevel(logging.INFO)

        sh = logging.StreamHandler(sys.stdout)
        sh.setFormatter(logging.Formatter(
            '%(asctime)s [%(levelname)s] %(message)s'))
        logger.addHandler(sh)

        # Create progress bar with ETA
        pbar = tqdm(desc="Translating", unit="chunk", dynamic_ncols=True)

        def cli_progress(current: int, total: int, status: str):
            if total > 0:
                pbar.total = total
                pbar.update(max(0, current - pbar.n))
                # Update postfix with status
                if "ETA:" in status:
                    eta = status.split("ETA:")[-1].strip()
                    pbar.set_postfix_str(f"ETA: {eta}")
                elif "Complete" in status:
                    pbar.set_postfix_str("✅ Done")
            else:
                pbar.update(1)

        def cli_log(msg: str):
            pass  # Handled by StreamHandler

        stop_event = threading.Event()

        try:
            stats = run_translation_task(
                input_path=args.input,
                output_path=args.output,
                model=args.model,
                api_type=args.api_type,
                direction=args.direction,
                output_format=args.format,
                api_key=api_key,
                ollama_url=ollama_url,
                progress_callback=cli_progress,
                log_callback=cli_log,
                stop_event=stop_event
            )
            pbar.close()

            print("\n" + "=" * 50)
            print("📊 TRANSLATION STATISTICS")
            print("=" * 50)
            for k, v in stats.items():
                if k not in ["total_segments", "total_chunks"]:
                    print(f"{k.replace('_', ' ').title()}: {v}")
            print(f"Total Segments: {stats.get('total_segments', 0)}")
            print(f"Total Chunks: {stats.get('total_chunks', 0)}")
            print("=" * 50)

        except KeyboardInterrupt:
            stop_event.set()
            pbar.close()
            print("\n⚠️ Translation interrupted by user.")
            sys.exit(1)
        except Exception as e:
            pbar.close()
            print(f"\n❌ Fatal Error: {e}")
            if args.verbose:
                import traceback
                traceback.print_exc()
            sys.exit(1)


if __name__ == "__main__":
    main()
